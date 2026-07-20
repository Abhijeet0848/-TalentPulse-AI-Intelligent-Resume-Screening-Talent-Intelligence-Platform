import os
import re
import json
from pypdf import PdfReader
from docx import Document

# Skills Dictionary Taxonomy
SKILLS_TAXONOMY = {
    # Programming Languages
    "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Golang", "Rust", 
    "PHP", "Ruby", "Swift", "Kotlin", "R", "Scala", "SQL", "Bash", "Shell", "HTML", "CSS", "Sass",
    
    # ML & AI / Data Science
    "Machine Learning", "Deep Learning", "Data Science", "Artificial Intelligence", "NLP",
    "Natural Language Processing", "Computer Vision", "TensorFlow", "PyTorch", "Keras",
    "Scikit-Learn", "Pandas", "NumPy", "Matplotlib", "Seaborn", "Plotly", "SciPy",
    "OpenCV", "Hugging Face", "Transformers", "BERT", "LLM", "Prompt Engineering",
    "Data Analysis", "Data Visualization", "Feature Engineering", "Model Deployment",
    
    # Frameworks & Libraries
    "React", "React.js", "Vue", "Vue.js", "Angular", "Node.js", "Express", "Next.js", 
    "Django", "Flask", "FastAPI", "Spring Boot", "ASP.NET", "Tailwind CSS", "Bootstrap",
    
    # Cloud & DevOps
    "AWS", "Amazon Web Services", "Azure", "GCP", "Google Cloud", "Docker", "Kubernetes",
    "CI/CD", "GitHub Actions", "Jenkins", "Terraform", "Ansible", "Linux", "Unix",
    "Microservices", "REST API", "GraphQL", "Serverless",
    
    # Databases & Big Data
    "PostgreSQL", "MySQL", "MongoDB", "Redis", "SQLite", "Oracle", "DynamoDB",
    "Apache Spark", "Hadoop", "Kafka", "Snowflake", "Elasticsearch", "ETL", "SQL Server",
    
    # Tools & Methods
    "Git", "GitHub", "GitLab", "Jira", "Confluence", "Agile", "Scrum", "TDD",
    "System Design", "OOP", "Object-Oriented Programming", "Unit Testing",
    
    # Management & Soft Skills
    "Project Management", "Product Management", "Leadership", "Team Collaboration",
    "Problem Solving", "Communication", "Critical Thinking", "Stakeholder Management",
    "Tableau", "Power BI", "Excel"
}

def extract_text_from_pdf(file_path):
    """Extract text content from PDF file."""
    text = ""
    try:
        reader = PdfReader(file_path)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        print(f"Error reading PDF {file_path}: {e}")
    return text

def extract_text_from_docx(file_path):
    """Extract text content from DOCX file."""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = " ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    text += row_text + "\n"
    except Exception as e:
        print(f"Error reading DOCX {file_path}: {e}")
    return text

def extract_text_from_txt(file_path):
    """Extract text content from plain text file."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception as e:
        print(f"Error reading TXT {file_path}: {e}")
        return ""

def parse_resume_file(file_path):
    """Detect file extension and extract raw text."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return extract_text_from_docx(file_path)
    else:
        return extract_text_from_txt(file_path)

def extract_email(text):
    """Extract email address using regex."""
    email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
    match = re.search(email_pattern, text)
    return match.group(0) if match else ""

def extract_phone(text):
    """Extract phone number using regex."""
    phone_pattern = r'(\+?\d{1,3}[\s-]?)?\(?\d{3}\)?[\s-]?\d{3}[\s-]?\d{4}'
    match = re.search(phone_pattern, text)
    return match.group(0) if match else ""

def extract_links(text):
    """Extract LinkedIn and GitHub links."""
    linkedin = ""
    github = ""
    
    linkedin_match = re.search(r'(https?://(?:www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+)', text, re.IGNORECASE)
    if linkedin_match:
        linkedin = linkedin_match.group(0)
        
    github_match = re.search(r'(https?://(?:www\.)?github\.com/[a-zA-Z0-9_-]+)', text, re.IGNORECASE)
    if github_match:
        github = github_match.group(0)
        
    return linkedin, github

def extract_candidate_name(text, file_name=""):
    """Extract candidate name from first few lines or fallback to filename."""
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    if lines:
        for line in lines[:5]:
            # Look for line that looks like a name (2-3 words, no numbers, no special symbols)
            if re.match(r'^[A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,2}$', line):
                return line
    # Fallback to filename without extension
    clean_fn = os.path.splitext(file_name)[0].replace("_", " ").replace("-", " ")
    return clean_fn.title() if clean_fn else "Unknown Candidate"

def extract_education(text):
    """Detect candidate's highest education degree level."""
    text_lower = text.lower()
    if re.search(r'\b(phd|doctorate|ph\.d)\b', text_lower):
        return "Ph.D."
    elif re.search(r'\b(master|masters|m\.s|m\.tech|mba|msc)\b', text_lower):
        return "Master's"
    elif re.search(r'\b(bachelor|bachelors|b\.s|b\.tech|bsc|ba|b\.e)\b', text_lower):
        return "Bachelor's"
    elif re.search(r'\b(associate|diploma)\b', text_lower):
        return "Associate/Diploma"
    return "Bachelor's"

def extract_experience_years(text):
    """Extract total years of experience using regex patterns."""
    text_lower = text.lower()
    
    # Direct mention: "5+ years of experience", "3 years experience", "4.5 years in software"
    exp_matches = re.findall(r'(\d+(?:\.\d+)?)\+?\s*(?:-\s*\d+)?\s*(?:years?|yrs?)\b(?:\s+of)?\s*(?:experience|exp)?', text_lower)
    if exp_matches:
        try:
            years = [float(m) for m in exp_matches if float(m) < 40]
            if years:
                return max(years)
        except ValueError:
            pass

    # Look for year ranges like "2018 - 2023" or "2020 - Present"
    year_ranges = re.findall(r'(20\d{2})\s*[-–—to]+\s*(20\d{2}|present|current)', text_lower)
    total_years = 0.0
    current_year = 2026
    for start_yr, end_yr in year_ranges:
        try:
            start = int(start_yr)
            end = current_year if end_yr in ["present", "current"] else int(end_yr)
            if 1990 <= start <= current_year and start <= end:
                total_years += (end - start)
        except ValueError:
            pass
            
    return round(min(total_years, 35.0), 1) if total_years > 0 else 2.0

def extract_skills(text):
    """Extract recognized skills from text matching skills taxonomy."""
    found_skills = set()
    text_lower = text.lower()
    
    for skill in SKILLS_TAXONOMY:
        # Regex boundary check to avoid partial word matches (e.g., "R" in "React")
        pattern = r'(?<![a-zA-Z0-9])' + re.escape(skill.lower()) + r'(?![a-zA-Z0-9])'
        if re.search(pattern, text_lower):
            found_skills.add(skill)
            
    return sorted(list(found_skills))

def parse_full_resume(file_path):
    """Run full parsing pipeline on a resume file."""
    raw_text = parse_resume_file(file_path)
    file_name = os.path.basename(file_path)
    
    name = extract_candidate_name(raw_text, file_name)
    email = extract_email(raw_text)
    phone = extract_phone(raw_text)
    linkedin, github = extract_links(raw_text)
    education = extract_education(raw_text)
    exp_years = extract_experience_years(raw_text)
    skills = extract_skills(raw_text)
    
    return {
        "full_name": name,
        "email": email,
        "phone": phone,
        "linkedin": linkedin,
        "github": github,
        "education_level": education,
        "experience_years": exp_years,
        "parsed_skills": skills,
        "raw_text": raw_text,
        "file_name": file_name,
        "file_path": file_path
    }
